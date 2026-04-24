# src/utils/exporters.py
import os
from docx import Document
from htmldocx import HtmlToDocx
from xhtml2pdf import pisa
from markdownify import markdownify as md

# Carpeta temporal para generar las exportaciones antes de enviarlas al usuario
EXPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'exports_temp')
os.makedirs(EXPORT_DIR, exist_ok=True)

def export_document(doc_data, formato):
    """Transforma el HTML de la BD al formato solicitado preservando estilos y autoría."""
    
    # Limpiamos el título para generar un nombre de archivo seguro para Windows/Linux
    safe_title = "".join([c for c in doc_data['titulo'] if c.isalpha() or c.isdigit() or c==' ']).rstrip()
    filepath = os.path.join(EXPORT_DIR, f"{safe_title}.{formato}")
    
    # Extraemos los datos crudos
    html_raw = doc_data['contenido_texto']
    autor_name = doc_data.get('autor') or "Autor no especificado"

    if formato == 'pdf':
        # Plantilla HTML inyectable para xhtml2pdf con estilos profesionales
        pdf_html = f"""
        <html>
        <head>
            <style>
                @page {{ size: A4; margin: 2.5cm 2cm; }}
                body {{ font-family: Helvetica, Arial, sans-serif; font-size: 11pt; line-height: 1.6; color: #333; }}
                h1.title {{ text-align: center; color: #1e293b; font-size: 24pt; margin-bottom: 5px; }}
                p.author {{ text-align: center; font-style: italic; color: #64748b; margin-top: 0; margin-bottom: 30px; }}
                hr {{ border-top: 1px solid #cbd5e1; margin-bottom: 20px; }}
                img {{ max-width: 100%; height: auto; display: block; margin: 10px auto; }}
                table {{ border-collapse: collapse; width: 100%; margin-bottom: 15px; }}
                th, td {{ border: 1px solid #cbd5e1; padding: 10px; text-align: left; }}
                th {{ background-color: #f8fafc; font-weight: bold; }}
                pre {{ background-color: #f1f5f9; padding: 10px; border-radius: 5px; font-family: monospace; }}
            </style>
        </head>
        <body>
            <h1 class="title">{doc_data['titulo']}</h1>
            <p class="author">Desarrollado por: {autor_name}</p>
            <hr>
            <div>
                {html_raw}
            </div>
        </body>
        </html>
        """
        # Creación del archivo físico PDF
        with open(filepath, "w+b") as result_file:
            pisa.CreatePDF(pdf_html, dest=result_file)

    elif formato == 'docx':
        # Instanciamos python-docx
        document = Document()
        document.add_heading(doc_data['titulo'], 0)
        document.add_paragraph(f"Desarrollado por: {autor_name}", style='Subtitle')
        
        # Puente para traducir HTML a estilos de Word
        new_parser = HtmlToDocx()
        new_parser.add_html_to_document(html_raw, document)
        
        document.save(filepath)

    elif formato == 'md':
        # Convertimos HTML a Markdown (ideal para wikis, GitHub o documentación técnica nativa)
        md_content = md(html_raw, heading_style="ATX")
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"# {doc_data['titulo']}\n")
            f.write(f"**Autor:** {autor_name}\n\n---\n\n{md_content}")

    elif formato == 'txt':
        # Texto plano: usamos markdownify para limpiar etiquetas pero manteniendo cierta legibilidad estructural
        texto_limpio = md(html_raw, strip=['a', 'img']) 
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"TITULO: {doc_data['titulo']}\n")
            f.write(f"AUTOR: {autor_name}\n")
            f.write(f"========================================\n\n")
            f.write(f"{texto_limpio}")

    return filepath